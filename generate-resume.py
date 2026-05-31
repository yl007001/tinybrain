#!/usr/bin/env python3
"""Generate TinyBrain project experience Word document for resume."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ========== Page Setup ==========
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ========== Styles ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Title style
title_style = doc.styles['Title']
title_style.font.size = Pt(22)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
h1.font.name = '微软雅黑'
h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
h2.font.name = '微软雅黑'
h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
    return p

def add_normal(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    run.font.size = Pt(8)

# ========== Content ==========

# Title
title = doc.add_heading('TinyBrain — 个人 AI 知识引擎', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('项目经历 · 核心贡献者 · 全栈开发者')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
run.font.name = '微软雅黑'

doc.add_paragraph()  # spacing

# ========== 基本信息 ==========
doc.add_heading('一、项目基本信息', level=1)

add_normal('项目名称：TinyBrain — 个人 AI 知识引擎')
add_normal('项目定位：基于 Spring Boot 3.x + Spring Cloud Alibaba 的全链路个人知识管理及 AI 问答系统')
add_normal('项目规模：~4,500+ 行 Java 代码，~2,000+ 行前端代码，7 个 Maven 模块，覆盖完整后端技术栈')
add_normal('项目周期：2026 年 5 月 — 至今')
add_normal('开源协议：Apache 2.0')
add_normal('')

# ========== 技术栈 ==========
doc.add_heading('二、核心技术栈', level=1)

tech_data = [
    ('后端语言', 'Java 17（Records、Pattern Matching、Virtual Threads）'),
    ('核心框架', 'Spring Boot 3.2.5、Spring Cloud 2023.0.3、Spring Cloud Alibaba'),
    ('ORM 层', 'MyBatis-Plus 3.5.7（分页、自动填充、逻辑删除）'),
    ('数据库', 'MySQL 8（生产）、H2（开发测试）'),
    ('认证授权', 'JWT（jjwt 0.12.5）、Spring Security、BCrypt 加密'),
    ('AI 集成', 'DeepSeek / OpenAI 兼容 API、Ollama 本地模型'),
    ('向量存储', '内存存储 + JSON 文件持久化（可替换 ChromaDB/Milvus）'),
    ('API 网关', 'Spring Cloud Gateway（WebFlux 响应式）'),
    ('容错保护', 'Resilience4j（熔断器、重试、限流器）'),
    ('可观测性', 'Micrometer + Prometheus + Grafana + Zipkin 链路追踪'),
    ('日志体系', 'Logback + Logstash JSON 编码 + MDC traceId 关联'),
    ('接口文档', 'SpringDoc OpenAPI 2.6（Swagger UI 自动生成）'),
    ('前端', 'Vue 3 + TypeScript + Element Plus + Vite'),
    ('部署', 'Docker、Docker Compose（MySQL/ES/Redis/Prometheus/Grafana/Zipkin）'),
    ('CI/CD', 'GitHub Actions（编译→测试→打包→Docker 构建）'),
]

for label, value in tech_data:
    add_bullet(value, bold_prefix=f'{label}：')

doc.add_paragraph()

# ========== 核心功能 ==========
doc.add_heading('三、核心功能模块', level=1)

doc.add_heading('1. RAG 检索增强生成系统', level=2)
add_normal('RAG（Retrieval Augmented Generation）是本项目的核心亮点，将传统知识库管理与 AI 大模型相结合，解决了 LLM 知识截止和幻觉问题。')

add_bullet('文档分块策略：实现自适应分块算法，支持按段落边界分割，块大小可配（500 字符 + 100 字符重叠），保证语义完整性', bold_prefix='')
add_bullet('向量化：通过 LLM Embedding API 将文本块转为语义向量（默认 1024 维），兼容 DeepSeek / OpenAI / Ollama', bold_prefix='')
add_bullet('向量存储：采用 ConcurrentHashMap + 读写锁 + 文件持久化，重启不丢数据；支持余弦相似度 Top-K 检索', bold_prefix='')
add_bullet('RAG 流程：问题向量化 → 语义检索 → 上下文拼接 → LLM 增强生成 → 带来源标注的回答', bold_prefix='')
add_bullet('监控埋点：使用 Micrometer @Timed 注解对索引和问答耗时进行监控', bold_prefix='')

doc.add_heading('2. Agent 智能体系统', level=2)
add_normal('基于 Function Calling 模式的 AI Agent，支持多工具自动调度和多轮对话记忆。')

add_bullet('Agent 引擎：支持工具注册、Function Calling 循环调度、最大迭代轮次控制（防无限循环）', bold_prefix='')
add_bullet('工具插件体系：接口 + Spring 自动注册机制，新增工具只需实现 AgentTool 接口 + @Component', bold_prefix='')
add_bullet('内置工具：知识库搜索（KnowledgeSearchTool，已接入 RAGService 实现真实检索）、计算器（CalculatorTool）、日期时间（DateTimeTool）、网络搜索（WebSearchTool）', bold_prefix='')
add_bullet('对话记忆：基于 ConcurrentHashMap 的 session 级上下文管理，保留最近 10 轮对话，支持会话重置', bold_prefix='')

doc.add_heading('3. 认证与权限系统', level=2)
add_bullet('JWT 认证：基于 jjwt 0.12.x，密钥通过环境变量 TINYBRAIN_JWT_SECRET 注入，开发环境自动兜底', bold_prefix='')
add_bullet('Spring Security 集成：OncePerRequestFilter + 自定义 UserDetailsService + BCrypt 密码加密', bold_prefix='')
add_bullet('RBAC 权限：@PreAuthorize 注解 + 角色判断，支持 ADMIN/USER 分级', bold_prefix='')

doc.add_heading('4. 知识库管理', level=2)
add_bullet('文档 CRUD：支持创建、更新、逻辑删除、分页查询（带关键词搜索 + 状态筛选）', bold_prefix='')
add_bullet('多类型支持：Markdown / 纯文本，标签系统（JSON 数组）', bold_prefix='')
add_bullet('文件上传：支持 .md / .txt 文件上传，自动提取标题和摘要', bold_prefix='')

doc.add_heading('5. 微服务与可观测性', level=2)
add_bullet('Spring Cloud Gateway 网关：JWT 全局鉴权 Filter、路由转发、CORS 跨域', bold_prefix='')
add_bullet('Docker Compose：一键启动 MySQL 8 + ES 7 + Redis 7 + Prometheus + Grafana + Zipkin + 应用', bold_prefix='')
add_bullet('CI/CD：GitHub Actions 流水线，含后端编译/测试/打包、Docker 镜像构建、前端构建', bold_prefix='')
add_bullet('生产配置：隔离的开发/生产配置文件、Resilience4j 熔断限流、日志轮转', bold_prefix='')

doc.add_paragraph()

# ========== 架构亮点 ==========
doc.add_heading('四、架构设计与亮点', level=1)

add_bullet('分模块 Maven 架构：7 个模块（common/user/knowledge/rag/agent/gateway/app），清晰的分层和依赖方向', bold_prefix='模块化设计：')
add_bullet('统一响应格式 R<T> + 全局异常处理（BusinessException + @RestControllerAdvice），涵盖参数校验、业务异常、系统异常', bold_prefix='统一异常处理：')
add_bullet('核心算法使用并行流（parallelStream）加速余弦相似度计算，读写锁（ReentrantReadWriteLock）保证并发安全', bold_prefix='并发安全：')
add_bullet('Resilience4j CircuitBreaker + Retry + RateLimiter 防止 LLM API 故障级联，VectorStore 延时批量写入减少 I/O', bold_prefix='容错与保护：')
add_bullet('向量存储 5s 延时写入 + PreDestroy 优雅关闭，保证数据不丢；应用健康检查（HealthIndicator）', bold_prefix='数据持久化：')
add_bullet('MDC traceId 贯穿 HTTP 请求 → 业务逻辑 → 外部 API 调用，Zipkin 分布式追踪可视化调用链', bold_prefix='全链路追踪：')
add_bullet('Docker 多阶段构建（构建镜像 200MB → 运行镜像 100MB+）、健康检查、JVM GC 调优参数', bold_prefix='容器化优化：')

doc.add_paragraph()

# ========== 面试考点 ==========
doc.add_heading('五、面试高频考点覆盖', level=1)

add_normal('该项目在设计时系统性地覆盖了以下面试常见考点，每个关键模块均有对应的代码实现和面试话术：')

topics = [
    ('Java 17 新特性', 'Records（DTO）、Pattern Matching、Virtual Threads（后续可启用）、Sealed Classes'),
    ('Spring Boot 自动配置', '@SpringBootApplication 组合注解、@Conditional、自动配置原理讲解文档'),
    ('AOP 与事务', '@Transactional 声明式事务、@Timed 自定义注解、Transaction 传播机制'),
    ('MyBatis-Plus', '分页插件、自动填充（MetaObjectHandler）、逻辑删除、条件构造器 LambdaQueryWrapper'),
    ('JWT 认证流程', 'Token 签发/解析/验证、Filter 链、SecurityContextHolder 线程上下文'),
    ('RAG 原理与优化', '分块策略、向量化模型选择、余弦相似度 vs 欧氏距离、检索质量评估'),
    ('Function Calling', 'Tool Use 模式原理、工具 Schema 定义、多轮循环控制、ReAct 模式'),
    ('微服务架构', '网关模式（Gateway）、服务发现（Nacos）、配置中心、分布式问题'),
    ('Docker 与 CI/CD', '多阶段构建、health check、GitHub Actions 流水线设计'),
    ('可观测性三支柱', 'Metrics（Prometheus）、Tracing（Zipkin）、Logging（ELK）'),
    ('容错设计', '熔断器（Circuit Breaker）原理、重试策略、限流、降级'),
    ('数据库优化', '索引设计、事务隔离级别、连接池（HikariCP）调优、SQL 优化'),
]

for topic, desc in topics:
    add_bullet(desc, bold_prefix=f'{topic}：')

doc.add_paragraph()

# ========== 工程实践 ==========
doc.add_heading('六、软件工程实践', level=1)

add_bullet('统一代码风格，JavaDoc 类注释 + 关键方法注释 + 面试考点说明', bold_prefix='代码质量：')
add_bullet('56+ 单元测试覆盖工具类、核心引擎（VectorStore、DocChunkStrategy、AgentEngine、JwtUtil 等）', bold_prefix='测试覆盖：')
add_bullet('双语 README（中英文）、CONTRIBUTING.md、CODE_OF_CONDUCT.md、Issue/PR 模板', bold_prefix='文档规范：')
add_bullet('Apache 2.0 开源协议、清晰的 Git 提交信息（conventional commits 风格）', bold_prefix='开源规范：')

doc.add_paragraph()

# ========== 个人角色 ==========
doc.add_heading('七、个人职责与贡献', level=1)

add_bullet('独立完成项目从 0 到 1 的架构设计、编码实现、测试部署', bold_prefix='')
add_bullet('设计并实现 RAG 检索增强生成系统（分块 → 向量化 → 语义检索 → LLM 生成）', bold_prefix='')
add_bullet('设计并实现 Agent 智能体系统（Function Calling 引擎 + 工具插件体系）', bold_prefix='')
add_bullet('搭建 Spring Cloud Gateway + Nacos 微服务架构', bold_prefix='')
add_bullet('配置 Prometheus + Grafana + Zipkin 可观测性体系', bold_prefix='')
add_bullet('编写 Docker Compose + GitHub Actions CI/CD 部署流水线', bold_prefix='')
add_bullet('编写 10 篇教学文档（从 Spring Boot 原理到面试包装）', bold_prefix='')
add_bullet('构建 Vue 3 + Element Plus 前端界面', bold_prefix='')

doc.add_paragraph()

# ========== 总结 ==========
doc.add_heading('八、项目总结', level=1)

add_normal(
    'TinyBrain 是一个覆盖 Java 后端全链路的技术项目，从 Spring Boot 基础框架一直延伸到 AI 大模型集成。'
    '它在设计时不仅关注功能实现，还系统性地考虑了面试场景——每个关键模块都有详细的面试考点说明，'
    '帮助面试者深入理解背后的原理和设计取舍。'
)
add_normal('')
add_normal(
    '项目最大的差异化优势在于：① RAG + Agent 双 AI 引擎的整合；② 全链路技术栈覆盖；'
    '③ 面试导向的设计文档。这使得它既是一个真实可用的知识库工具，也是求职面试的强力加分项。'
)

# ========== Save ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TinyBrain-项目经历.docx')
doc.save(output_path)
print(f'Resume document saved to: {output_path}')
