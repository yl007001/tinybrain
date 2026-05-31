#!/usr/bin/env python3
"""Generate comprehensive job hunting strategy document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# Page Setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

title_style = doc.styles['Title']
title_style.font.size = Pt(22)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

h1 = doc.styles['Heading 1']
h1.font.size = Pt(15)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
h1.font.name = '微软雅黑'
h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

h2 = doc.styles['Heading 2']
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
h2.font.name = '微软雅黑'
h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

h3 = doc.styles['Heading 3']
h3.font.size = Pt(11)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
h3.font.name = '微软雅黑'
h3.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_normal(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    return p

def add_bullet(text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
    return p

def add_table_with_data(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()

# ============================================================
# CONTENT
# ============================================================

title = doc.add_heading('TinyBrain 求职战略指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('针对：民办本科 · 无实习 · 目标 Java 后端 大厂 25K-30K')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()

# ========== 第一章 ==========
doc.add_heading('一、背景分析与定位策略', level=1)

doc.add_heading('1.1 你的定位', level=2)
add_normal('学历背景：民办本科（非 985/211），无大厂实习经历')
add_normal('项目优势：TinyBrain 全链路 AI 知识引擎，覆盖 7 个模块，技术栈全面且有深度（IVF、RAG、Agent、微服务）')
add_normal('核心策略：用项目深度突破学历壁垒。面试官看到 TinyBrain 的 IVF 索引实现、Resilience4j 容错体系、SSE 流式架构时，' +
           '会认可你的工程能力。目标是让面试官说出："这不像一个应届生写的代码。"')

doc.add_heading('1.2 薪资定位', level=2)
add_table_with_data(
    ['档位', '月薪', '目标公司档次', '可行性', '策略'],
    [
        ['第一档', '30K+', '阿里/字节/腾讯/美团', '难度极大', '简历关极难过，除非项目+博客+开源 star 加持'],
        ['第二档', '22K-28K', '快手/拼多多/京东/百度/滴滴/小米', '有机会', '主攻方向，项目够深可弥补学历'],
        ['第三档', '18K-22K', '中型互联网/独角兽/二线大厂', '稳妥', '保底方向，面试压力较小'],
        ['第四档', '15K-18K', '小型公司/外包/创业公司', '轻松', '完全不建议，你的项目水平远超此档'],
    ]
)

add_normal('建议策略：主攻第二档（22K-28K），冲刺第一档（30K+），用第三档保底。')

doc.add_paragraph()

# ========== 第二章 ==========
doc.add_heading('二、目标公司与岗位清单', level=1)

doc.add_heading('2.1 第一梯队（冲刺）', level=2)
add_table_with_data(
    ['公司', '岗位名称', '城市', '预计薪资', '投递渠道', '备注'],
    [
        ['字节跳动', '后端开发工程师-抖音/电商/飞书', '北京/上海/深圳', '28K-35K', '官网/内推', '算法要求高，项目要能讲出深度'],
        ['阿里巴巴', 'Java 开发工程师-业务平台/本地生活', '杭州/北京', '25K-32K', '官网/内推/Boss', '学历卡最严，简历关最难过'],
        ['腾讯', '后台开发-云/广告/微信', '深圳/北京', '25K-32K', '官网/内推', 'Linux/C++/Go 都可能问'],
        ['美团', '后端开发-到家/到店/优选', '北京/上海', '24K-30K', '官网/Boss', '业务面要求高，问得细'],
    ]
)

doc.add_heading('2.2 第二梯队（主攻）', level=2)
add_table_with_data(
    ['公司', '岗位名称', '城市', '预计薪资', '投递渠道', '备注'],
    [
        ['快手', 'Java 后端开发', '北京', '24K-30K', '官网/Boss/内推', '面试难度低于字节阿里'],
        ['拼多多', '服务端研发工程师', '上海', '25K-32K', '官网', '薪资高但 11116，应届生慎重'],
        ['京东', 'Java 开发工程师', '北京', '22K-28K', '官网/Boss', '技术氛围相对宽松'],
        ['百度', 'Java 后端研发', '北京', '22K-28K', '官网', 'AI 方向加分，TinyBrain 对口'],
        ['滴滴', '后端研发工程师', '北京', '22K-28K', '官网/Boss', '业务稳定性高'],
        ['小米', 'Java 服务端开发', '北京/南京', '20K-26K', '官网/Boss', '学历要求相对友好'],
        ['B站', 'Java 后端开发', '上海', '22K-28K', '官网/Boss', '技术氛围好，面试有深度'],
        ['网易', 'Java 开发工程师', '杭州/广州', '22K-27K', '官网', '面试重基础和项目'],
        ['得物/小红书', '后端开发', '上海', '20K-28K', 'Boss/内推', '增长期，招人多'],
    ]
)

doc.add_heading('2.3 第三梯队（保底）', level=2)
add_table_with_data(
    ['公司', '岗位名称', '城市', '预计薪资', '投递渠道', '备注'],
    [
        ['携程', 'Java 后端开发', '上海', '18K-24K', '官网/Boss', '工作生活平衡'],
        ['贝壳找房', '后端研发', '北京', '18K-24K', '官网/Boss', '技术栈匹配'],
        ['58同城', 'Java 开发', '北京', '16K-22K', 'Boss', '面试较简单'],
        ['跟谁学/好未来', '后端开发', '北京', '18K-24K', '官网/Boss', '教育行业稳定'],
        ['趣头条/一点资讯', '后端开发', '北京/上海', '16K-22K', 'Boss', '小厂练手用'],
    ]
)

doc.add_paragraph()

# ========== 第三章 ==========
doc.add_heading('三、投递渠道与策略', level=1)

doc.add_heading('3.1 渠道优先级', level=2)
add_table_with_data(
    ['渠道', '优先级', '优势', '劣势', '操作建议'],
    [
        ['牛客网内推', 'S级', '可避开简历筛选直接进面试', '需要找到靠谱内推人', '每天刷牛客找内推帖，主动私信'],
        ['学长学姐内推', 'S级', '成功率最高', '需要有人脉', '联系同校学长，即使不认识也礼貌请教'],
        ['Boss直聘', 'A级', '回复快，可聊薪资', '信息杂乱', '完善在线简历+作品集链接'],
        ['公司官网投递', 'A级', '官方渠道，校招必走', '简历筛选严', '校招季必须每个公司都投'],
        ['脉脉', 'B级', '可直连HR/技术Leader', '需要身份认证', '关注目标公司招聘动态'],
        ['拉勾/猎聘', 'B级', '中高级岗位多', '应届生岗位少', '补充渠道'],
    ]
)

doc.add_heading('3.2 简历投递策略', level=2)
add_normal('核心原则：海投 + 精准。')

add_bullet('海投：目标公司清单中的所有公司，不管能不能进，先投了再说', bold_prefix='')
add_bullet('精准：对每一家公司，针对其业务特点微调简历（比如投美团就强调高并发经验）', bold_prefix='')
add_bullet('数量目标：截至毕业前，至少投递 200+ 份简历，100+ 次笔试，30+ 次面试', bold_prefix='')
add_bullet('投递节奏：每周集中投 10-15 家，保持持续的输出', bold_prefix='')

doc.add_heading('3.3 内推话术', level=2)
add_normal('给不认识的人发内推请求时：')
add_normal('"学长/学姐你好，我是 XX 学校的 26 届毕业生，主攻 Java 后端。我独立做了一个 AI 知识引擎项目（GitHub 开源），' +
           '技术栈是 Spring Boot 3 + RAG + Agent，想投贵司的 Java 后端岗位。这是我的简历和项目链接，方便帮忙内推吗？谢谢！"')
add_normal('要点：礼貌、简洁、附上简历+项目链接、说明为什么想要这个机会。')

doc.add_paragraph()

# ========== 第四章 ==========
doc.add_heading('四、简历优化指南', level=1)

doc.add_heading('4.1 针对 TinyBrain 的简历写法', level=2)
add_normal('标准写法（已为你生成在 TinyBrain-项目经历.docx 中）：')
add_bullet('项目名称+时间+技术栈（一行看清楚）', bold_prefix='→ ')
add_bullet('3-4 个 bullet point，每个包含：技术名词 + 做了什么 + 量化成果', bold_prefix='→ ')
add_bullet('不要写"熟悉 Java"，要写"在 TinyBrain 中使用 Java 17 Records/Pattern Matching 优化 DTO 代码"', bold_prefix='→ ')

doc.add_heading('4.2 针对不同公司的微调', level=2)
add_table_with_data(
    ['投递公司', '简历强调点'],
    [
        ['字节跳动', '高并发：IVF 索引加速 50x + Resilience4j 容错 + @Async 异步索引'],
        ['阿里巴巴', '微服务：Spring Cloud Gateway + Nacos + Docker Compose 全栈部署'],
        ['腾讯', 'Linux/网络：Docker 多阶段构建 + JMeter 压测 + Zipkin 链路追踪'],
        ['百度', 'AI 能力：RAG 检索增强 + Agent Function Calling + IVF ANN 索引'],
        ['快手/美团', '业务理解：知识库 CRUD + 分页优化 + 缓存策略 + 数据库索引设计'],
        ['小米/京东', '全栈能力：Vue3 前端 + Spring Boot 后端 + Docker 部署'],
    ]
)

doc.add_paragraph()

# ========== 第五章 ==========
doc.add_heading('五、面试准备路线图', level=1)

doc.add_heading('5.1 八股文（必须滚瓜烂熟）', level=2)
add_table_with_data(
    ['领域', '必问题', 'TinyBrain 对应', '优先级'],
    [
        ['Java', '集合/并发/JVM/反射', '项目中使用 ConcurrentHashMap/并行流/读写锁', '★★★★★'],
        ['Spring', 'IoC/AOP/事务/循环依赖', '项目大量使用 @Transactional/@Timed/声明式事务', '★★★★★'],
        ['MySQL', '索引/隔离级别/锁/Explain', '项目设计索引/分页/逻辑删除', '★★★★★'],
        ['Redis', '数据结构/缓存/穿透/雪崩', '项目配置 Redis 缓存 RAG 结果', '★★★★☆'],
        ['网络', 'TCP/IP/HTTP/HTTPS', '项目中 WebClient 调用 LLM API', '★★★★☆'],
        ['操作系统', '进程线程/内存管理', '项目中线程池调优 @Async', '★★★☆☆'],
        ['设计模式', '策略/工厂/单例/代理', '项目使用策略模式(AgentTool)/模板方法', '★★★☆☆'],
        ['分布式', 'CAP/一致性/服务发现', '项目配置 Nacos/Gateway', '★★★☆☆'],
    ]
)

doc.add_heading('5.2 项目深挖（TinyBrain 专有）', level=2)
add_normal('以下问题必须能用中文+英文各讲一遍：')
add_bullet('RAG 的完整流程是什么？你做了哪些优化？（查询改写 + IVF + 批量嵌入 + 融合搜索）', bold_prefix='')
add_bullet('IVF 索引的原理？K 和 nprobe 怎么取值？（K≈√N, nprobe≈2%K, K-means++初始化）', bold_prefix='')
add_bullet('Resilience4j 的三层保护怎么配置的？（熔断3次10s+重试2次1s+限流20次/分）', bold_prefix='')
add_bullet('SSE 和 WebSocket 有什么区别？为什么 RAG 用 SSE？（单向推送、天然 HTTP、简单）', bold_prefix='')
add_bullet('向量搜索 O(n) 到 O(√n) 怎么实现的？召回率多少？（IVF 聚类分桶，召回~90%+暴力融合）', bold_prefix='')
add_bullet('如果知识库有一百万篇文档，系统会怎样？怎么优化？（分库+ES+kafka 异步索引）', bold_prefix='')

doc.add_heading('5.3 算法题准备', level=2)
add_normal('大厂面试必考算法。建议刷题计划：')
add_table_with_data(
    ['阶段', '周期', '题量', '重点'],
    [
        ['基础期', '第1-2周', '50题', '数组/链表/栈/队列/哈希表（LeetCode Hot 100）'],
        ['进阶期', '第3-4周', '50题', '二叉树/动态规划/贪心/回溯'],
        ['冲刺期', '第5-6周', '50题', '高频题二刷+公司题库（牛客搜索目标公司面经）'],
        ['维持期', '直到入职', '每天2题', '保持手感，不熟悉的题重做'],
    ]
)
add_normal('推荐顺序：LeetCode 热题 100 → 剑指 Offer → 目标公司题库（牛客网面经汇总）')

doc.add_paragraph()

# ========== 第六章 ==========
doc.add_heading('六、补充加分项', level=1)

doc.add_heading('6.1 技术博客（强烈推荐）', level=2)
add_normal('写 3-5 篇技术文章发布到掘金/知乎/CSDN/V2EX：')
add_bullet('"从 O(n) 到 O(√n)：在 Java 中实现 IVF 倒排索引做 ANN 搜索"', bold_prefix='推荐主题1：')
add_bullet('"RAG 检索增强生成从入门到生产：查询改写 + 混合搜索 + 重排序"', bold_prefix='推荐主题2：')
add_bullet('"面试官问烂了的 JWT，我从源码角度给你讲透"', bold_prefix='推荐主题3：')
add_bullet('"应届生做了个 AI 知识引擎，面试官说这代码不像应届生写的"', bold_prefix='推荐主题4：')
add_normal('发表后转发到牛客网「技术交流」板块，附带项目链接。')

doc.add_heading('6.2 GitHub 开源运营', level=2)
add_normal('目标：100+ stars')
add_bullet('README 加截图（Swagger 界面、前端界面、Grafana 大盘）', bold_prefix='')
add_bullet('发到 HackerNews / 掘金 / V2EX / 开源中国', bold_prefix='')
add_bullet('在项目的 Issues 里记录 TODO 和改进计划，展示你是 actively maintained', bold_prefix='')
add_bullet('提交到 awesome-spring-boot / awesome-rag 等 curated list', bold_prefix='')

doc.add_heading('6.3 在线 Demo（加分极大）', level=2)
add_normal('买一台云服务器（阿里云轻量应用服务器 2C4G 约 68元/月）：')
add_bullet('docker-compose up -d 一键部署', bold_prefix='')
add_bullet('配置 Nginx 反代 + 域名（如 tinybrain.xxx.com）', bold_prefix='')
add_bullet('简历和博客里直接附上线 Demo 链接', bold_prefix='')
add_normal('面试官看到能直接访问的 Demo，印象分会极大提升。')

doc.add_paragraph()

# ========== 第七章 ==========
doc.add_heading('七、投递时间线', level=1)

add_table_with_data(
    ['时间', '事项', '具体行动'],
    [
        ['5月-6月', '项目打磨+面试准备', '完善 TinyBrain、刷题、写博客、部署 Demo'],
        ['6月-7月', '提前批投递', '投递字节/阿里/腾讯/百度/快手的提前批（hc 最多）'],
        ['7月-8月', '秋招正式批', '海投所有目标公司，每周至少投递 15-20 家'],
        ['8月-10月', '面试高峰期', '面试、复盘、继续投递并行'],
        ['10月-11月', '收割 Offer', '已有 Offer 的可以谈薪资，没有的继续面'],
        ['11月-12月', '补录+春招预备', '如果还没满意 Offer，准备明年春招'],
        ['次年2月-4月', '春招', '最后机会，此时竞争更激烈但 hc 还有'],
    ]
)

add_normal('关键节点：6-7月提前批是最好时机，hc 最多且竞争相对较小。一定不要错过。')

doc.add_paragraph()

# ========== 第八章 ==========
doc.add_heading('八、薪资谈判指南', level=1)

doc.add_heading('8.1 谈薪策略', level=2)
add_bullet('永远不要主动说期望薪资，让 HR 先说', bold_prefix='第1条：')
add_bullet('手里有 Offer 才有底气谈薪，所以面试阶段要分散排期', bold_prefix='第2条：')
add_bullet('拿到 Offer 后用 A 公司的 Offer 去聊 B 公司的涨幅', bold_prefix='第3条：')
add_bullet('对应届生来说，公积金比例比月薪差一两千更重要', bold_prefix='第4条：')

doc.add_heading('8.2 薪资构成解析', level=2)
add_table_with_data(
    ['公司', '月薪base', '年终奖', '股票/期权', '公积金', '实际年薪估算'],
    [
        ['字节', 'base 按 12 月发', '3-6 个月', '有期权', '12% 全额', 'base×(15-18)'],
        ['阿里', 'base 按 12 月发', '3-6 个月', '有股票', '12% 全额', 'base×(15-18)+股票'],
        ['腾讯', 'base 按 12 月发', '2-6 个月', '有股票', '12% 全额', 'base×(14-18)+股票'],
        ['美团', 'base 按 12 月发', '2-4 个月', '有期权', '12% 部分', 'base×(14-16)'],
    ]
)
add_normal('注意：不要只看月薪，要把年终奖、公积金、股票全算进去。例如：月薪 25K + 年终 4 个月 = 25×16=40万/年。')

doc.add_paragraph()

# ========== 第九章 ==========
doc.add_heading('九、常见踩坑点', level=1)

add_bullet('等着不投 → 等准备好了再投，结果错过了黄金时间', bold_prefix='坑1：')
add_bullet('只投一家 → 所有鸡蛋放一个篮子，挂了就没了', bold_prefix='坑2：')
add_bullet('简历写太长 → 一页就够了，面试官没时间看两页', bold_prefix='坑3：')
add_bullet('八股文背了但不会用 → 面试官问"你项目里怎么用的"答不上来', bold_prefix='坑4：')
add_bullet('算法刷太少 → 大厂一票否决，200 题是底线', bold_prefix='坑5：')
add_bullet('不说项目亮点 → 面试官问项目就简单说"我做了个 XX"，不会包装', bold_prefix='坑6：')
add_bullet('不谈薪资 → 被 HR 压价，少拿好几千', bold_prefix='坑7：')

doc.add_paragraph()

# ========== 第十章 ==========
doc.add_heading('十、一句话总结', level=1)

add_normal('你的核心优势不是学历，不是实习，而是「一个人写了一个大厂 3 人团队才做得出来的项目」。把这个故事讲好，就能突破学历天花板。')
add_normal('')
add_normal('行动清单：')
add_bullet('本周：完善简历、写第一篇博客、部署 Demo', bold_prefix='')
add_bullet('本月：刷 50 题、投递 20 家公司、准备项目深挖话术', bold_prefix='')
add_bullet('下月：持续投递 + 面试 + 复盘，循环到拿到 Offer', bold_prefix='')

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TinyBrain-求职战略指南.docx')
doc.save(output_path)
print(f'Job strategy guide saved to: {output_path}')
