#!/usr/bin/env python3
"""Generate TinyBrain user guide document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

title_style = doc.styles['Title']
title_style.font.size = Pt(22)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

h1 = doc.styles['Heading 1']
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
h1.font.name = '微软雅黑'
h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

h2 = doc.styles['Heading 2']
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
h2.font.name = '微软雅黑'
h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_normal(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    return p

def add_table_with_data(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    doc.add_paragraph()

# ============================================================
title = doc.add_heading('TinyBrain 完全使用指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('个人 AI 知识引擎 —— Spring Boot 3.x + RAG + Agent')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

doc.add_paragraph()

# ========== 第一章 ==========
doc.add_heading('一、项目是什么', level=1)

add_normal('TinyBrain 是一个个人 AI 知识引擎。你可以把它理解为「你自己的知识库 + AI 问答系统」。')
add_normal('')
add_normal('能做什么：')
add_bullet('上传文档到知识库（支持 Markdown/普通文本文件）', bold_prefix='📚 ')
add_bullet('对文档执行「索引」操作，自动分块并转成向量存入搜索引擎', bold_prefix='🔍 ')
add_bullet('像 ChatGPT 一样提问，但回答基于你自己的文档内容', bold_prefix='💬 ')
add_bullet('Agent 智能体模式：可以调用工具（计算器、查时间、搜知识库）', bold_prefix='🤖 ')
add_normal('')
add_normal('适合谁用：')
add_bullet('个人：管理笔记、技术文档，用 AI 快速检索', bold_prefix='')
add_bullet('学生：把教材/笔记导进去，直接问问题', bold_prefix='')
add_bullet('面试者：作为简历项目展示全栈能力', bold_prefix='')

doc.add_paragraph()

# ========== 第二章 ==========
doc.add_heading('二、项目结构速览', level=1)

add_normal('项目在 tinybrain/ 目录下，核心结构：')
add_code('tinybrain/')
add_code('├── tinybrain-app/        # 🔵 启动入口（从这里启动）')
add_code('├── tinybrain-common/     # 公共模块（配置、异常、工具类）')
add_code('├── tinybrain-user/      # 用户模块（登录注册、JWT 认证）')
add_code('├── tinybrain-knowledge/ # 知识库模块（文档 CRUD）')
add_code('├── tinybrain-rag/       # 🔵 RAG 核心（向量检索、LLM 调用）')
add_code('├── tinybrain-agent/     # 🔵 Agent 模块（智能体、工具插件）')
add_code('├── tinybrain-gateway/   # API 网关（路由转发、鉴权）')
add_code('├── tinybrain-ui/        # 🖥️ Vue3 前端界面')
add_code('├── deploy/             # Docker 部署配置（Prometheus/Grafana）')
add_code('└── docs/               # 📖 文档（面试资料、使用指南）')

doc.add_paragraph()

# ========== 第三章 ==========
doc.add_heading('三、数据库在哪里？', level=1)

add_normal('这是新手最容易迷惑的地方。TinyBrain 的数据库取决于你用什么「模式」启动：')
doc.add_paragraph()

add_table_with_data(
    ['启动模式', '数据库类型', '数据存储位置', '是否持久化', '适用场景'],
    [
        ['dev（开发模式）', 'H2 内存数据库', '仅在内存中', '❌ 重启丢数据', '开发调试、学习'],
        ['docker（部署模式）', 'MySQL 8', 'Docker Volume', '✅ 永久保存', '正式使用、Demo'],
        ['prod（生产模式）', 'MySQL 8', '本地/云数据库', '✅ 永久保存', '生产环境'],
    ]
)

doc.add_heading('3.1 开发模式（H2 内存数据库）', level=2)
add_normal('配置位置：tinybrain-app/src/main/resources/application-dev.yml')
add_normal('当使用 dev 模式启动时，数据库是 H2 内存数据库：')
add_bullet('类型：H2（纯 Java 内存数据库，无需安装）', bold_prefix='')
add_bullet('数据存在：内存中（JVM 退出即消失）', bold_prefix='')
add_bullet('URL：jdbc:h2:mem:tinybrain', bold_prefix='')
add_bullet('控制台：http://localhost:8080/h2-console（账号 sa，密码空）', bold_prefix='')
add_normal('表结构：启动时自动执行 schema.sql 初始化（位置：tinybrain-app/src/main/resources/sql/schema.sql）')
add_normal('这就是为什么你在电脑上找不到 MySQL 数据库——因为 dev 模式根本不需要 MySQL，它运行在内存里。')

doc.add_heading('3.2 Docker 模式（MySQL 8）', level=2)
add_normal('当使用 docker-compose 启动时：')
add_bullet('MySQL 服务运行在 Docker 容器中，容器名：tinybrain-mysql', bold_prefix='')
add_bullet('数据持久化在 Docker Volume：mysql-data', bold_prefix='')
add_bullet('数据库名：tinybrain，账号 root，密码 tinybrain2026', bold_prefix='')
add_bullet('端口映射到宿主机 3306，可以用 Navicat/DataGrip 连接', bold_prefix='')
add_normal('Docker 模式下连接 MySQL：')
add_code('Host: localhost')
add_code('Port: 3306')
add_code('User: root')
add_code('Password: tinybrain2026')
add_code('Database: tinybrain')

doc.add_heading('3.3 向量数据存在哪里？', level=2)
add_normal('除了普通数据库，TinyBrain 还有一个「向量数据库」（存文档的数学向量）：')
add_bullet('开发模式下存在内存 + JSON 文件：./data/vectorstore/vectors.json', bold_prefix='')
add_bullet('Docker 模式下存在 Docker Volume：app-data', bold_prefix='')
add_normal('这意味着即使重启应用，向量数据也不会丢（JSON 文件持久化）。')

doc.add_paragraph()

# ========== 第四章 ==========
doc.add_heading('四、如何启动项目', level=1)

doc.add_heading('4.1 前提条件', level=2)
add_bullet('JDK 17+（必须）', bold_prefix='')
add_bullet('Maven 3.9+（必须）', bold_prefix='')
add_bullet('Git（可选）', bold_prefix='')
add_bullet('Docker & Docker Compose（Docker 模式需要）', bold_prefix='')
add_bullet('Node.js（前端需要）', bold_prefix='')
add_normal('验证环境：')
add_code('java -version        # 确认 Java 17')
add_code('mvn -version        # 确认 Maven')
add_code('docker compose version  # 确认 Docker（可选）')

doc.add_heading('4.2 快速启动（开发模式，推荐第一次尝试）', level=2)
add_normal('这种模式不需要安装 MySQL，不需要 Docker，开箱即用：')
add_code('# 1. 进入项目目录')
add_code('cd tinybrain')
add_code('')
add_code('# 2. 编译（第一次需要下载依赖，几分钟）')
add_code('mvn clean install -DskipTests')
add_code('')
add_code('# 3. 启动后端')
add_code('cd tinybrain-app')
add_code('mvn spring-boot:run -Dspring-boot.run.profiles=dev')
add_normal('启动后访问 http://localhost:8080/swagger-ui.html 可以看到 API 文档。')

doc.add_heading('4.3 启动前端（可选）', level=2)
add_normal('新开一个终端窗口：')
add_code('cd tinybrain/tinybrain-ui')
add_code('npm install          # 安装依赖')
add_code('npm run dev         # 启动前端')
add_normal('启动后访问 http://localhost:3000，界面会自动代理到后端 8080 端口。')

doc.add_heading('4.4 Docker 模式启动', level=2)
add_normal('这种模式会启动 MySQL + Redis + 应用 + Prometheus + Grafana + Zipkin 全套服务：')
add_code('# 在项目根目录')
add_code('docker-compose up -d')
add_code('')
add_code('# 查看启动日志')
add_code('docker-compose logs -f tinybrain-app')
add_normal('启动大概需要 1-2 分钟，因为要初始化数据库。')

doc.add_paragraph()

# ========== 第五章 ==========
doc.add_heading('五、LLM API Key 配置', level=1)

add_normal('TinyBrain 的 AI 能力需要调用大模型 API。你需要一个 API Key（二选一）：')

add_table_with_data(
    ['平台', '注册地址', '费用', '推荐场景'],
    [
        ['DeepSeek', 'platform.deepseek.com', '注册送 500 万 token', '首选，便宜量大'],
        ['OpenAI', 'platform.openai.com', '付费', '备选，需海外信用卡'],
        ['Ollama（本地）', 'ollama.ai', '免费', '完全离线，不需要 Key'],
    ]
)

add_normal('配置方式（二选一）：')
add_normal('方式一：设置环境变量（推荐）')
add_code('# Windows PowerShell')
add_code('$env:TINYBRAIN_LLM_KEY="sk-your-key-here"')
add_code('')
add_code('# Linux/Mac')
add_code('export TINYBRAIN_LLM_KEY=sk-your-key-here')
add_code('')
add_code('# 然后启动应用即可')
add_code('mvn spring-boot:run -Dspring-boot.run.profiles=dev')

add_normal('方式二：修改配置文件 application.yml')
add_code('# 修改 tinybrain-app/src/main/resources/application.yml')
add_code('tinybrain:')
add_code('  llm:')
add_code('    api-key: sk-your-key-here   # 改成你的 Key')

add_normal('如果不想申请 API Key，可以用 Ollama 本地模式：')
add_code('# 1. 安装 Ollama：https://ollama.ai')
add_code('# 2. 拉取模型')
add_code('ollama pull qwen2.5:7b')
add_code('ollama pull nomic-embed-text')
add_code('')
add_code('# 3. 启动应用（使用 Ollama profile）')
add_code('cd tinybrain-app')
add_code('mvn spring-boot:run -Dspring-boot.run.profiles=ollama')
add_normal('注意：Ollama 模式下不需要 API Key，但需要本地有 GPU 或足够的内存。')

doc.add_paragraph()

# ========== 第六章 ==========
doc.add_heading('六、第一次使用流程', level=1)

add_normal('假设你已经启动了后端，可以按以下步骤体验完整功能：')

doc.add_heading('Step 1: 注册账号', level=2)
add_code('# 用 curl 或 Swagger UI 注册')
add_code('curl -X POST http://localhost:8080/api/auth/register \\')
add_code('  -H "Content-Type: application/json" \\')
add_code("  -d '{\"username\":\"test\",\"password\":\"123456\"}'")
add_normal('或者打开 Swagger UI（http://localhost:8080/swagger-ui.html），找到 Auth 模块的 /register 接口。')

doc.add_heading('Step 2: 登录获取 Token', level=2)
add_code('curl -X POST http://localhost:8080/api/auth/login \\')
add_code('  -H "Content-Type: application/json" \\')
add_code("  -d '{\"username\":\"test\",\"password\":\"123456\"}'")
add_normal('返回结果中的 token 需要记下来，后续所有操作都需要它。')

doc.add_heading('Step 3: 创建一篇文档', level=2)
add_code('# 用刚才拿到的 TOKEN 替换下面的 xxx')
add_code('curl -X POST http://localhost:8080/api/documents \\')
add_code('  -H "Content-Type: application/json" \\')
add_code('  -H "Authorization: Bearer xxx" \\')
add_code("  -d '{\"title\":\"Spring 事务\",\"content\":\"Spring 事务的传播机制有 REQUIRED、REQUIRES_NEW、NESTED 等。REQUIRED 是默认的，如果当前有事务则加入，没有则新建。\"}'")

doc.add_heading('Step 4: 索引到 RAG（关键步骤）', level=2)
add_code('curl -X POST "http://localhost:8080/api/rag/index/1" \\')
add_code('  -H "Authorization: Bearer xxx"')
add_normal('这一步会把文档分块 → 向量化 → 存入向量库。如果不做这步，后续问答查不到内容。')

doc.add_heading('Step 5: 开始 RAG 问答', level=2)
add_code('curl "http://localhost:8080/api/rag/ask?question=Spring事务有哪些传播机制&topK=3" \\')
add_code('  -H "Authorization: Bearer xxx"')
add_normal('如果 API Key 配置正确，会返回基于文档内容的 AI 回答。')

doc.add_paragraph()

# ========== 第七章 ==========
doc.add_heading('七、前端界面使用', level=1)

add_normal('启动前端（npm run dev）后，打开 http://localhost:3000：')
add_bullet('登录/注册页面 → 注册后自动跳转控制台', bold_prefix='1️⃣ ')
add_bullet('控制台 → 查看系统状态和快速入口', bold_prefix='2️⃣ ')
add_bullet('知识库 → 上传 .md/.txt 文件或手动创建文档', bold_prefix='3️⃣ ')
add_bullet('RAG 问答 → 输入问题，AI 基于你的文档回答', bold_prefix='4️⃣ ')
add_bullet('AI Agent → 智能体对话，可调用计算器、查时间等工具', bold_prefix='5️⃣ ')

doc.add_paragraph()

# ========== 第八章 ==========
doc.add_heading('八、常见问题与排错', level=1)

add_table_with_data(
    ['问题', '原因', '解决'],
    [
        ['启动报错 "端口 8080 被占用"', '其他程序在用 8080', '关掉其他程序，或在 application.yml 改 server.port'],
        ['登录返回 401', 'Token 过期或未正确设置', '重新登录获取新 Token，确认 Authorization 头格式为 Bearer xxx'],
        ['RAG 问答返回 "未检索到相关信息"', '文档未索引 / 向量库为空', '先 POST /api/rag/index/{id} 索引文档'],
        ['RAG 问答返回 "LLM API 未配置"', '没设置 API Key', '设置环境变量 TINYBRAIN_LLM_KEY，或用 Ollama 模式'],
        ['创建文档返回 403', '没有 Token', '先登录获取 Token，加到请求头'],
        ['前端页面白屏', '没启动前端 / 端口不对', '确认 npm run dev 运行中，访问 http://localhost:3000'],
        ['Docker 启动 MySQL 报错', '3306 端口被本地 MySQL 占用', '先关掉本地 MySQL，或改 docker-compose 端口映射'],
        ['H2 控制台打不开', 'dev 模式没启用', '确认启动参数是 --spring.profiles.active=dev'],
    ]
)

doc.add_paragraph()

# ========== 第九章 ==========
doc.add_heading('九、访问地址速查表', level=1)

add_table_with_data(
    ['服务', '地址', '说明'],
    [
        ['API 文档 (Swagger)', 'http://localhost:8080/swagger-ui.html', '所有接口的在线文档和测试工具'],
        ['前端界面', 'http://localhost:3000', 'Vue 3 图形界面'],
        ['H2 数据库控制台', 'http://localhost:8080/h2-console', 'dev 模式下查看数据库（账号 sa，密码空）'],
        ['健康检查', 'http://localhost:8080/actuator/health', '查看应用运行状态'],
        ['Prometheus 指标', 'http://localhost:9090', 'Docker 模式下的监控系统'],
        ['Grafana 大盘', 'http://localhost:3000', 'Docker 模式下的可视化面板（admin/admin）'],
        ['Zipkin 链路追踪', 'http://localhost:9411', 'Docker 模式下的请求追踪'],
    ]
)

doc.add_paragraph()

# ========== 第十章 ==========
doc.add_heading('十、关键命令速查', level=1)

add_code('# 编译项目')
add_code('mvn clean install -DskipTests')
add_code('')
add_code('# 启动后端 (开发模式, H2 数据库)')
add_code('cd tinybrain-app && mvn spring-boot:run -Dspring-boot.run.profiles=dev')
add_code('')
add_code('# 启动后端 (Ollama 本地模式)')
add_code('cd tinybrain-app && mvn spring-boot:run -Dspring-boot.run.profiles=ollama')
add_code('')
add_code('# 启动前端')
add_code('cd tinybrain-ui && npm run dev')
add_code('')
add_code('# 运行全部测试')
add_code('mvn test')
add_code('')
add_code('# Docker 一键部署')
add_code('docker-compose up -d')
add_code('')
add_code('# 一键脚本启动')
add_code('./deploy.sh dev       # 开发模式')
add_code('./deploy.sh docker   # Docker 模式')
add_code('./deploy.sh ollama   # 本地模型模式')

doc.add_paragraph()
add_normal('至此 TinyBrain 的基础使用你已经掌握了。有任何具体问题可以随时问我。')

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TinyBrain-使用指南.docx')
doc.save(output_path)
print(f'User guide saved to: {output_path}')
